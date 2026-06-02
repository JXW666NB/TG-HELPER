# -*- coding: utf-8 -*-
"""
TG Helper - 工具集 (全面升级版)
包含：文件操作、系统控制、网络爬虫、浏览器自动化（反爬增强）、
      AI 图像/视频分析、物联网设备管理、QQ 扩展、多媒体编辑、
      Arduino 嵌入式开发、定时任务、VPN/代理工具等。
"""
import os
import sys
import shutil
import subprocess
import base64
import webbrowser
import pyautogui
import time
import re
import requests
import random
import asyncio
from datetime import datetime
from openai import OpenAI
from config import config
from skill_manager import SkillManager


def _deepseek_sdk_params(model_name):
    if not model_name or 'deepseek' not in str(model_name).lower():
        return {}
    params = {}
    if getattr(config, 'deepseek_thinking_enabled', False):
        params['reasoning_effort'] = getattr(config, 'deepseek_reasoning_effort', 'high')
        params['extra_body'] = {"thinking": {"type": "enabled"}}
    ctx = getattr(config, 'deepseek_context_window', 0)
    if ctx:
        params['max_tokens'] = ctx
    return params


class Tools:
    def __init__(self, memory, confirm_callback=None, output_callback=None, task_scheduler=None, skill_manager=None, gui=None):
        self.memory = memory
        self.confirm_callback = confirm_callback or (lambda msg: True)
        self.output_callback = output_callback or (lambda msg: None)
        self.task_scheduler = task_scheduler
        self.skill_manager = skill_manager or SkillManager(getattr(config, 'skills_dirs', ["./skills"]))
        self.gui = gui
        self.client = OpenAI(
            api_key=config.ai_api_key,
            base_url=config.ai_base_url
        )
        # 插件注册的工具表：{tool_name: {"description": str, "parameters": dict, "plugin_id": str}}
        self._plugin_tools: dict = {}
        os.makedirs("./screenshots", exist_ok=True)
        os.makedirs("./downloads", exist_ok=True)

    def register_plugin_tool(self, tool_name: str, description: str, parameters: dict, plugin_id: str):
        """供插件系统调用，注册插件提供的工具"""
        self._plugin_tools[tool_name] = {
            "description": description,
            "parameters": parameters,
            "plugin_id": plugin_id,
        }

    def unregister_plugin_tool(self, tool_name: str):
        """卸载插件工具"""
        self._plugin_tools.pop(tool_name, None)

    def get_plugin_tools_summary(self) -> str:
        """生成插件工具的文本摘要，供 AI 系统提示词使用"""
        if not self._plugin_tools:
            return ""
        lines = ["\n【插件注册的工具】"]
        for name, info in self._plugin_tools.items():
            desc = info.get("description", "")
            params = info.get("parameters", {})
            param_str = ""
            if isinstance(params, dict) and "properties" in params:
                props = params["properties"]
                req = params.get("required", [])
                parts = []
                for k, v in props.items():
                    mark = "(必填)" if k in req else "(可选)"
                    parts.append(f"{k}{mark}")
                param_str = ", ".join(parts)
            elif isinstance(params, dict):
                param_str = ", ".join(params.keys())
            lines.append(f"- {name}({param_str}) - {desc}")
        return "\n".join(lines)

    def _expand_path(self, path):
        """展开路径中的 ~ 并返回绝对路径"""
        if path is None:
            return None
        return os.path.abspath(os.path.expanduser(str(path)))

    # ==================== 基础系统命令 ====================
    def execute_command(self, command: str, cwd: str = None):
        """执行系统命令，返回输出（编码安全版，兼容Windows GBK）"""
        try:
            if cwd:
                cwd = self._expand_path(cwd)
            env = os.environ.copy()
            env['PYTHONIOENCODING'] = 'utf-8'
            env['PYTHONUTF8'] = '1'
            result = subprocess.run(
                command,
                shell=True,
                cwd=cwd,
                capture_output=True,
                timeout=30,
                env=env
            )

            def _safe_decode(data: bytes) -> str:
                if not data:
                    return ""
                for encoding in ('utf-8', 'gbk', 'gb2312', 'cp936', 'latin-1'):
                    try:
                        return data.decode(encoding)
                    except (UnicodeDecodeError, LookupError):
                        continue
                return data.decode('utf-8', errors='replace')

            return {
                "stdout": _safe_decode(result.stdout),
                "stderr": _safe_decode(result.stderr),
                "returncode": result.returncode
            }
        except Exception as e:
            return {"error": str(e)}

    def read_file(self, filepath: str, max_chars: int = 8000000):
        """读取文件内容，支持 .txt, .docx, .ino"""
        filepath = self._expand_path(filepath)
        ext = os.path.splitext(filepath)[1].lower()
        try:
            if ext == '.txt':
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif ext == '.docx':
                try:
                    from docx import Document
                except ImportError:
                    return "ERROR: Missing required library 'python-docx'. Please use 'install_python_package' to install it."
                doc = Document(filepath)
                full_text = [para.text for para in doc.paragraphs]
                content = '\n'.join(full_text)
            elif ext == '.ino':
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                return f"ERROR: Unsupported file format: {ext}"

            total_len = len(content)
            if total_len == 0:
                return "INFO: The file is empty."
            if total_len <= max_chars:
                return f"SUCCESS: File read successfully. Content ({total_len} characters):\n{content}"
            else:
                preview = content[:max_chars]
                return f"INFO: File is large ({total_len} characters). Showing first {max_chars} characters:\n{preview}\n\nTo read more, use 'read_file_chunk' with start={max_chars}, chunk_size={max_chars}."
        except FileNotFoundError:
            return f"ERROR: File not found: {filepath}. Please check the path."
        except Exception as e:
            return f"ERROR: {str(e)}"

    def read_file_chunk(self, filepath: str, start: int = 0, chunk_size: int = 9000000):
        """分块读取文件内容"""
        filepath = self._expand_path(filepath)
        ext = os.path.splitext(filepath)[1].lower()
        try:
            if ext == '.txt':
                with open(filepath, 'r', encoding='utf-8') as f:
                    full_content = f.read()
            elif ext == '.docx':
                try:
                    from docx import Document
                except ImportError:
                    return "ERROR: Missing required library 'python-docx'."
                doc = Document(filepath)
                full_text = [para.text for para in doc.paragraphs]
                full_content = '\n'.join(full_text)
            else:
                return f"ERROR: Unsupported format: {ext}"
            total_len = len(full_content)
            end = min(start + chunk_size, total_len)
            chunk = full_content[start:end]
            has_more = end < total_len
            if has_more:
                return f"INFO: Chunk {start}-{end}/{total_len}:\n{chunk}\n\nTo continue, use 'read_file_chunk' with start={end}, chunk_size={chunk_size}."
            else:
                return f"SUCCESS: Final chunk {start}-{end}/{total_len}:\n{chunk}\n\nFile fully read."
        except Exception as e:
            return f"ERROR: {str(e)}"

    def write_file(self, filepath: str, content: str):
        """写入文件"""
        filepath = self._expand_path(filepath)
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            return "SUCCESS: File written successfully."
        except Exception as e:
            return f"ERROR: Failed to write file: {str(e)}"

    def delete_file(self, filepath: str):
        """删除文件（可配置确认）"""
        filepath = self._expand_path(filepath)
        if config.tool_confirmation.get("delete_file", True):
            confirmed = self.confirm_callback(f"确认删除 {filepath}？")
            if not confirmed:
                return "INFO: Deletion cancelled by user."
        try:
            os.remove(filepath)
            return "SUCCESS: File deleted successfully."
        except Exception as e:
            return f"ERROR: Failed to delete file: {str(e)}"

    def list_directory(self, path: str = "."):
        """列出目录内容"""
        path = self._expand_path(path)
        try:
            items = os.listdir(path)
            return f"SUCCESS: Directory listing for '{path}':\n" + "\n".join(items)
        except FileNotFoundError:
            return f"ERROR: Directory not found: {path}."
        except Exception as e:
            return f"ERROR: Failed to list directory: {str(e)}"

    def open_browser(self, url: str):
        """用系统默认浏览器打开URL"""
        webbrowser.open(url)
        return f"SUCCESS: Browser opened to {url}"

    def screenshot(self):
        """截取整个屏幕"""
        img = pyautogui.screenshot()
        path = f"./screenshots/screenshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        img.save(path)
        return f"SUCCESS: Screenshot saved to {path}"

    def read_tool_prompt(self, keyword: str):
        """
        按关键词查找并读取工具提示词文件（推荐代替 read_file 读取 tool_prompts）。
        只需传简短关键词，自动匹配完整文件名。例如：
        - keyword="网络" → 匹配"网络与下载还有浏览器自动化网页操作(...).txt"
        - keyword="QQ" → 匹配"QQ交互(...).txt"
        - keyword="桌面" → 匹配"桌面程序自动化(...).txt"
        """
        prompt_dir = "./tool_prompts"
        if not os.path.isdir(prompt_dir):
            return "ERROR: tool_prompts 目录不存在"
        keyword_lower = keyword.lower().strip()
        for fname in os.listdir(prompt_dir):
            if keyword_lower in fname.lower():
                filepath = os.path.join(prompt_dir, fname)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                    return f"SUCCESS: 已读取 {fname}\n\n{content}"
                except Exception as e:
                    return f"ERROR: 读取失败: {e}"
        # 没找到，列出所有可用文件
        files = [f for f in os.listdir(prompt_dir) if f.endswith('.txt')]
        return f"ERROR: 未找到包含 '{keyword}' 的工具文件。\n可用文件列表:\n" + "\n".join(f"  - {f}" for f in files)

    def install_python_package(self, package_name: str):
        """安装Python包（可配置确认）"""
        if config.tool_confirmation.get("install_python_package", True):
            confirmed = self.confirm_callback(f"确认安装Python包 {package_name}？")
            if not confirmed:
                return "INFO: 安装已取消。"
        result = self.execute_command(f"pip install {package_name}")
        if result["returncode"] == 0:
            return f"SUCCESS: Package '{package_name}' installed successfully."
        else:
            return f"ERROR: Failed to install '{package_name}': {result['stderr']}"

    # ==================== 文件操作增强 ====================
    def move_file(self, source: str, destination: str):
        """移动或重命名文件/文件夹（可配置确认）"""
        source = self._expand_path(source)
        destination = self._expand_path(destination)
        conflict_msg = ""
        if os.path.exists(destination):
            conflict_msg = f"目标文件已存在，将被覆盖。"
        if config.tool_confirmation.get("move_file", True):
            confirmed = self.confirm_callback(f"确认移动文件？\n源: {source}\n目标: {destination}\n{conflict_msg}")
            if not confirmed:
                return "INFO: Move operation cancelled by user."
        try:
            dest_dir = os.path.dirname(destination)
            if dest_dir:
                os.makedirs(dest_dir, exist_ok=True)
            shutil.move(source, destination)
            return f"SUCCESS: Moved '{source}' to '{destination}'."
        except FileNotFoundError:
            return f"ERROR: Source not found: {source}"
        except Exception as e:
            return f"ERROR: Failed to move: {str(e)}"

    def copy_file(self, source: str, destination: str):
        """复制文件"""
        source = self._expand_path(source)
        destination = self._expand_path(destination)
        conflict_msg = ""
        if os.path.exists(destination):
            conflict_msg = f"目标文件已存在，将被覆盖。"
        if config.tool_confirmation.get("copy_file", True):
            confirmed = self.confirm_callback(f"确认复制文件？\n源: {source}\n目标: {destination}\n{conflict_msg}")
            if not confirmed:
                return "INFO: Copy operation cancelled by user."
        try:
            dest_dir = os.path.dirname(destination)
            if dest_dir:
                os.makedirs(dest_dir, exist_ok=True)
            shutil.copy2(source, destination)
            return f"SUCCESS: Copied '{source}' to '{destination}'."
        except FileNotFoundError:
            return f"ERROR: Source not found: {source}"
        except Exception as e:
            return f"ERROR: Failed to copy: {str(e)}"

    def create_directory(self, path: str):
        """创建文件夹（可递归）"""
        path = self._expand_path(path)
        try:
            os.makedirs(path, exist_ok=True)
            return f"SUCCESS: Directory created at '{path}'."
        except Exception as e:
            return f"ERROR: Failed to create directory: {str(e)}"

    def get_file_info(self, path: str):
        """获取文件/文件夹信息"""
        path = self._expand_path(path)
        try:
            stat = os.stat(path)
            info = {
                "path": path,
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "is_dir": os.path.isdir(path),
                "is_file": os.path.isfile(path),
            }
            return f"SUCCESS: File info:\n" + "\n".join(f"{k}: {v}" for k, v in info.items())
        except FileNotFoundError:
            return f"ERROR: Path not found: {path}"
        except Exception as e:
            return f"ERROR: Failed to get info: {str(e)}"

    # ==================== 系统信息 ====================
    def system_info(self):
        """获取CPU、内存、磁盘信息"""
        try:
            import psutil
        except ImportError:
            return "ERROR: Missing required library 'psutil'. Please install it using 'pip install psutil'."
        try:
            cpu_percent = psutil.cpu_percent(interval=1)
            memory = psutil.virtual_memory()
            disk = psutil.disk_usage('/')
            info = f"""
CPU使用率: {cpu_percent}%
内存总量: {memory.total / (1024**3):.2f} GB, 可用: {memory.available / (1024**3):.2f} GB, 使用率: {memory.percent}%
磁盘C: 总容量: {disk.total / (1024**3):.2f} GB, 已用: {disk.used / (1024**3):.2f} GB, 可用: {disk.free / (1024**3):.2f} GB, 使用率: {disk.percent}%
操作系统: {os.name}, 平台: {sys.platform}
            """
            return f"SUCCESS: 系统信息:\n{info.strip()}"
        except Exception as e:
            return f"ERROR: 获取系统信息失败: {str(e)}"

    # ==================== 网络与下载 ====================
    def batch_download(self, urls: list, output_dir: str = "./downloads"):
        """批量下载视频/文件"""
        if not isinstance(urls, list):
            return "ERROR: urls 必须是一个列表。"
        results = []
        output_dir = self._expand_path(output_dir)
        os.makedirs(output_dir, exist_ok=True)
        try:
            import yt_dlp
        except ImportError:
            return "ERROR: Missing required library 'yt-dlp'. Please install it using 'pip install yt-dlp'."
        for url in urls:
            try:
                ydl_opts = {'outtmpl': os.path.join(output_dir, '%(title)s.%(ext)s')}
                with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                    ydl.download([url])
                results.append(f"成功下载: {url}")
            except Exception as e:
                results.append(f"下载失败 {url}: {str(e)}")
        return f"SUCCESS: 批量下载完成。\n" + "\n".join(results)

    def extract_web_content(self, url: str, selector: str = None, extract_type: str = "text", summarize: bool = True):
        """提取网页内容（静态）。extract_type=html 且 summarize=True 时自动 AI 提炼。"""
        try:
            import requests as req_lib
            from bs4 import BeautifulSoup
        except ImportError:
            return "ERROR: Missing required libraries 'requests' and 'beautifulsoup4'."
        try:
            response = req_lib.get(url, timeout=10)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            if selector:
                elements = soup.select(selector)
                if not elements:
                    return f"INFO: 未找到匹配选择器 '{selector}' 的元素。"
                if extract_type == "text":
                    content = "\n".join(el.get_text() for el in elements)
                else:
                    content = "\n".join(str(el) for el in elements)
            else:
                if extract_type == "text":
                    content = soup.get_text()
                else:
                    content = response.text
            if summarize and len(content) > 3000 and extract_type == "html":
                summary = self._summarize_with_ai(content, "HTML", f"目标网址: {url}。注意保留关键 URL 链接。")
                if summary:
                    return f"SUCCESS: 页面 AI 分析摘要 (原始 {len(content)} 字符):\n{summary}"
            return f"SUCCESS: 提取内容:\n{content[:2000]}" + ("..." if len(content) > 2000 else "")
        except Exception as e:
            return f"ERROR: 提取失败: {str(e)}"

    def web_search(self, query: str, num_results: int = 5):
        """Google Custom Search（需配置API）"""
        api_key = getattr(config, 'google_api_key', None)
        cse_id = getattr(config, 'google_cse_id', None)
        if not api_key or not cse_id:
            return "ERROR: 需要配置 Google API Key 和 CSE ID。"
        try:
            url = "https://www.googleapis.com/customsearch/v1"
            params = {'q': query, 'key': api_key, 'cx': cse_id, 'num': num_results}
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()
            items = data.get('items', [])
            if not items:
                return "INFO: 未找到搜索结果。"
            result_text = ""
            for item in items:
                result_text += f"标题: {item['title']}\n链接: {item['link']}\n摘要: {item.get('snippet', '')}\n\n"
            return f"SUCCESS: 搜索结果:\n{result_text}"
        except Exception as e:
            return f"ERROR: 搜索失败: {str(e)}"

    # ==================== AI 图片生成 ====================
    def generate_image(self, prompt: str, size: str = None, n: int = 1, reference_images=None):
        """AI 文生图/图生图：根据文本描述 + 可选参考图生成图片。
        
        size 支持: "1024x1024" 像素尺寸 或 "1:1"/"16:9"/"9:16"/"4:3"/"3:4" 等比例字符串。
        reference_images: 可选，参考图列表（URL、本地路径 或 base64），用于图生图。
        """
        try:
            from image_generator import generate_image as _gen
        except ImportError:
            return "ERROR: image_generator 模块未加载。"
        try:
            result = _gen(prompt, size=size, n=n, reference_images=reference_images)
            return result
        except Exception as e:
            return f"ERROR: 图片生成失败: {str(e)}"

    # ==================== 多模态分析 ====================
    def analyze_image(self, image_path: str, prompt: str = "请描述图片的内容。"):
        """使用多模态模型分析图片（优先备用模型）"""
        image_path = self._expand_path(image_path)
        try:
            with open(image_path, "rb") as f:
                image_data = f.read()
            ext = os.path.splitext(image_path)[1].lstrip('.')
            if ext.lower() in ['jpg', 'jpeg']:
                mime = 'jpeg'
            else:
                mime = ext.lower()
            base64_image = base64.b64encode(image_data).decode('utf-8')
            image_url = f"data:image/{mime};base64,{base64_image}"

            use_multimodal = (config.multimodal_enabled and
                              config.multimodal_api_key and
                              config.multimodal_base_url and
                              config.multimodal_model)
            if use_multimodal:
                client = OpenAI(api_key=config.multimodal_api_key, base_url=config.multimodal_base_url)
                model = config.multimodal_model
                temperature = config.multimodal_temperature
                max_tokens = config.multimodal_max_tokens
            else:
                client = self.client
                model = config.ai_model
                temperature = config.temperature
                max_tokens = config.max_tokens

            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是擅长理解图像和视频。"},
                    {
                        "role": "user",
                        "content": [
                            {"type": "image_url", "image_url": {"url": image_url}},
                            {"type": "text", "text": prompt}
                        ]
                    }
                ],
                temperature=temperature,
                max_tokens=max_tokens,
                **_deepseek_sdk_params(model),
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"ERROR: Image analysis failed: {str(e)}"

    def analyze_video(self, video_path: str, prompt: str = "请描述视频的内容。"):
        """分析本地视频（多模态）"""
        video_path = self._expand_path(video_path)
        try:
            with open(video_path, "rb") as f:
                video_data = f.read()
            ext = os.path.splitext(video_path)[1].lstrip('.')
            base64_video = base64.b64encode(video_data).decode('utf-8')
            video_url = f"data:video/{ext};base64,{base64_video}"

            use_multimodal = (config.multimodal_enabled and
                              config.multimodal_api_key and
                              config.multimodal_base_url and
                              config.multimodal_model)
            if use_multimodal:
                client = OpenAI(api_key=config.multimodal_api_key, base_url=config.multimodal_base_url)
                model = config.multimodal_model
                temperature = config.multimodal_temperature
                max_tokens = config.multimodal_max_tokens
            else:
                client = self.client
                model = config.ai_model
                temperature = config.temperature
                max_tokens = config.max_tokens

            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是擅长理解图像和视频。"},
                    {
                        "role": "user",
                        "content": [
                            {"type": "video_url", "video_url": {"url": video_url}},
                            {"type": "text", "text": prompt}
                        ]
                    }
                ],
                temperature=temperature,
                max_tokens=max_tokens,
                **_deepseek_sdk_params(model),
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"ERROR: Video analysis failed: {str(e)}"

    def download_video(self, url: str, output_dir: str = "./downloads"):
        """下载网络视频（yt-dlp）"""
        output_dir = self._expand_path(output_dir)
        os.makedirs(output_dir, exist_ok=True)
        try:
            import yt_dlp
        except ImportError:
            return "ERROR: yt-dlp not installed. Please use 'install_python_package' to install 'yt-dlp'."
        try:
            ydl_opts = {'outtmpl': os.path.join(output_dir, '%(title)s.%(ext)s')}
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
            return f"SUCCESS: Video downloaded to directory {output_dir}."
        except Exception as e:
            return f"ERROR: Download failed: {str(e)}"

    def analyze_web_video(self, url: str, prompt: str = "请描述视频的内容。"):
        """下载并分析网络视频（简化版）"""
        download_result = self.download_video(url)
        if download_result.startswith("ERROR"):
            return download_result
        return download_result + " 请使用 analyze_video 工具指定具体文件进行分析。"

    # ==================== 键盘鼠标模拟 ====================
    def click(self, x: int = None, y: int = None, button: str = 'left', clicks: int = 1):
        """鼠标点击（可配置确认）"""
        if config.tool_confirmation.get("click", True):
            confirmed = self.confirm_callback(f"确认执行鼠标点击操作？坐标: ({x}, {y}) 按钮: {button}")
            if not confirmed:
                return "INFO: 操作已取消。"
        try:
            if x is not None and y is not None:
                pyautogui.click(x, y, button=button, clicks=clicks)
            else:
                pyautogui.click(button=button, clicks=clicks)
            return f"SUCCESS: 鼠标点击完成。"
        except Exception as e:
            return f"ERROR: 鼠标点击失败: {str(e)}"

    def type_text(self, text: str, interval: float = 0.0):
        """模拟键盘输入文本"""
        if config.tool_confirmation.get("type_text", True):
            confirmed = self.confirm_callback(f"确认输入文本: {text[:50]}" + ("..." if len(text) > 50 else ""))
            if not confirmed:
                return "INFO: 操作已取消。"
        try:
            pyautogui.typewrite(text, interval=interval)
            return f"SUCCESS: 文本输入完成。"
        except Exception as e:
            return f"ERROR: 文本输入失败: {str(e)}"

    def move_mouse(self, x: int, y: int, duration: float = 0.5):
        """移动鼠标到指定坐标"""
        if config.tool_confirmation.get("move_mouse", True):
            confirmed = self.confirm_callback(f"确认移动鼠标到 ({x}, {y})？")
            if not confirmed:
                return "INFO: 操作已取消。"
        try:
            pyautogui.moveTo(x, y, duration=duration)
            return f"SUCCESS: 鼠标移动到 ({x}, {y})。"
        except Exception as e:
            return f"ERROR: 鼠标移动失败: {str(e)}"

    # ==================== 办公操作 ====================
    def read_excel(self, filepath: str, sheet_name: str = None):
        """读取Excel文件"""
        filepath = self._expand_path(filepath)
        try:
            import pandas as pd
        except ImportError:
            return "ERROR: Missing required library 'pandas'. Please install it using 'pip install pandas'."
        try:
            df = pd.read_excel(filepath, sheet_name=sheet_name) if sheet_name else pd.read_excel(filepath)
            preview = df.head(10).to_string()
            return f"SUCCESS: 读取成功，共 {len(df)} 行，预览如下:\n{preview}"
        except Exception as e:
            return f"ERROR: 读取Excel失败: {str(e)}"

    def write_excel(self, filepath: str, data: list, columns: list = None):
        """写入Excel文件"""
        filepath = self._expand_path(filepath)
        try:
            import pandas as pd
        except ImportError:
            return "ERROR: Missing required library 'pandas'."
        try:
            df = pd.DataFrame(data, columns=columns)
            df.to_excel(filepath, index=False)
            return f"SUCCESS: 数据已写入 {filepath}"
        except Exception as e:
            return f"ERROR: 写入Excel失败: {str(e)}"

    def send_email(self, to: str, subject: str, body: str, attachments: list = None):
        """发送邮件（需配置）"""
        smtp_server = getattr(config, 'email_smtp_server', None)
        smtp_port = getattr(config, 'email_port', 587)
        email_user = getattr(config, 'email_user', None)
        email_password = getattr(config, 'email_password', None)
        if not all([smtp_server, email_user, email_password]):
            return "ERROR: 邮件配置不完整，请在config中设置。"
        try:
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart
            from email.mime.base import MIMEBase
            from email import encoders
        except ImportError:
            return "ERROR: 邮件模块导入失败。"
        try:
            msg = MIMEMultipart()
            msg['From'] = email_user
            msg['To'] = to
            msg['Subject'] = subject
            msg.attach(MIMEText(body, 'plain'))
            if attachments:
                for filepath in attachments:
                    filepath = self._expand_path(filepath)
                    with open(filepath, 'rb') as f:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(f.read())
                        encoders.encode_base64(part)
                        part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(filepath)}"')
                        msg.attach(part)
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(email_user, email_password)
            server.send_message(msg)
            server.quit()
            return f"SUCCESS: 邮件已发送至 {to}"
        except Exception as e:
            return f"ERROR: 发送邮件失败: {str(e)}"

    def generate_pdf(self, filepath: str, title: str, content: str):
        """生成简易PDF"""
        filepath = self._expand_path(filepath)
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
        except ImportError:
            return "ERROR: Missing required library 'reportlab'. Please install it using 'pip install reportlab'."
        try:
            c = canvas.Canvas(filepath, pagesize=A4)
            width, height = A4
            c.setFont("Helvetica", 12)
            text_object = c.beginText(50, height - 50)
            text_object.setFont("Helvetica", 12)
            text_object.textLine(title)
            text_object.textLine("=" * len(title))
            for line in content.split('\n'):
                text_object.textLine(line)
            c.drawText(text_object)
            c.save()
            return f"SUCCESS: PDF已生成: {filepath}"
        except Exception as e:
            return f"ERROR: 生成PDF失败: {str(e)}"

    # ==================== 开发辅助 ====================
    def execute_code(self, code: str, language: str = "python"):
        """执行Python代码片段（危险操作，可配置确认）"""
        if config.tool_confirmation.get("execute_code", True):
            confirmed = self.confirm_callback(f"确认执行以下代码？\n{code[:200]}" + ("..." if len(code) > 200 else ""))
            if not confirmed:
                return "INFO: 代码执行已取消。"
        if language.lower() != "python":
            return f"ERROR: 目前仅支持Python，不支持 {language}。"
        import tempfile
        with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False, encoding='utf-8') as f:
            f.write(code)
            tmpfile = f.name
        try:
            env = os.environ.copy()
            env['PYTHONIOENCODING'] = 'utf-8'
            result = subprocess.run(['python', tmpfile], capture_output=True, text=True, timeout=30,
                                    encoding='utf-8', errors='replace', env=env)
            output = result.stdout + result.stderr
            return f"SUCCESS: 执行结果:\n{output[:2000]}" + ("..." if len(output) > 2000 else "")
        except subprocess.TimeoutExpired:
            return "ERROR: 执行超时。"
        except Exception as e:
            return f"ERROR: 执行失败: {str(e)}"
        finally:
            os.unlink(tmpfile)

    def execute_python(self, code: str):
        """AI 常用别名，等同于 execute_code（执行Python代码片段）"""
        return self.execute_code(code)

    def git_clone(self, repo_url: str, dest_dir: str = "."):
        """克隆Git仓库"""
        dest_dir = self._expand_path(dest_dir)
        try:
            import git
        except ImportError:
            return "ERROR: Missing required library 'GitPython'. Please install it using 'pip install GitPython'."
        try:
            git.Repo.clone_from(repo_url, dest_dir)
            return f"SUCCESS: 仓库克隆到 {dest_dir}"
        except Exception as e:
            return f"ERROR: 克隆失败: {str(e)}"

    def git_commit(self, repo_path: str, message: str):
        """Git提交"""
        repo_path = self._expand_path(repo_path)
        try:
            import git
        except ImportError:
            return "ERROR: Missing required library 'GitPython'."
        try:
            repo = git.Repo(repo_path)
            repo.git.add(A=True)
            repo.index.commit(message)
            return f"SUCCESS: 提交成功: {message}"
        except Exception as e:
            return f"ERROR: 提交失败: {str(e)}"

    def git_push(self, repo_path: str, remote: str = "origin", branch: str = "main"):
        """Git推送"""
        repo_path = self._expand_path(repo_path)
        try:
            import git
        except ImportError:
            return "ERROR: Missing required library 'GitPython'."
        try:
            repo = git.Repo(repo_path)
            repo.remotes[remote].push(branch)
            return f"SUCCESS: 推送到 {remote}/{branch} 成功"
        except Exception as e:
            return f"ERROR: 推送失败: {str(e)}"

    # ==================== 搜索与信息检索 ====================
    def search_files(self, directory: str, pattern: str, search_content: bool = False):
        """搜索文件名或内容"""
        directory = self._expand_path(directory)
        if not os.path.isdir(directory):
            return f"ERROR: 目录不存在: {directory}"
        if not pattern or not pattern.strip():
            return "ERROR: pattern 不能为空"
        if pattern.strip() in ('*', '.', '.*', '?', 'all'):
            pattern = '[^/]'
        try:
            matches = []
            compiled = re.compile(pattern, re.IGNORECASE)
            for root, dirs, files in os.walk(directory):
                for file in files:
                    filepath = os.path.join(root, file)
                    if compiled.search(file):
                        matches.append(filepath)
                    elif search_content:
                        try:
                            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                                if compiled.search(f.read(1024*10)):
                                    matches.append(filepath + " (内容匹配)")
                        except:
                            pass
            if matches:
                return f"SUCCESS: 找到 {len(matches)} 个匹配项:\n" + "\n".join(matches[:50])
            else:
                return f"INFO: 未找到匹配项。"
        except Exception as e:
            return f"ERROR: 搜索失败: {str(e)}"

    # ==================== 多媒体编辑 ====================
    def edit_image(self, image_path: str, output_path: str = None, operation: str = "resize", **kwargs):
        """图片编辑（resize, crop, rotate, grayscale, blur, sharpen）"""
        image_path = self._expand_path(image_path)
        if not output_path:
            base, ext = os.path.splitext(image_path)
            output_path = f"{base}_edited{ext}"
        output_path = self._expand_path(output_path)
        try:
            from PIL import Image, ImageFilter
        except ImportError:
            return "ERROR: Missing required library 'Pillow'. Please install it using 'pip install Pillow'."
        try:
            img = Image.open(image_path)
            if operation == "resize":
                width = kwargs.get('width')
                height = kwargs.get('height')
                if width and height:
                    img = img.resize((width, height))
                else:
                    return "ERROR: resize需要提供width和height参数。"
            elif operation == "crop":
                left = kwargs.get('left')
                top = kwargs.get('top')
                right = kwargs.get('right')
                bottom = kwargs.get('bottom')
                if all([left, top, right, bottom]):
                    img = img.crop((left, top, right, bottom))
                else:
                    return "ERROR: crop需要提供left,top,right,bottom参数。"
            elif operation == "rotate":
                degrees = kwargs.get('degrees', 90)
                img = img.rotate(degrees, expand=True)
            elif operation == "grayscale":
                img = img.convert('L')
            elif operation == "blur":
                img = img.filter(ImageFilter.BLUR)
            elif operation == "sharpen":
                img = img.filter(ImageFilter.SHARPEN)
            else:
                return f"ERROR: 不支持的操作 '{operation}'。"
            img.save(output_path)
            return f"SUCCESS: 图片已保存到 {output_path}"
        except Exception as e:
            return f"ERROR: 图片编辑失败: {str(e)}"

    def audio_process(self, input_path: str, output_path: str = None, operation: str = "convert", **kwargs):
        """音频处理"""
        input_path = self._expand_path(input_path)
        if not output_path:
            base, ext = os.path.splitext(input_path)
            output_path = f"{base}_processed{ext}"
        output_path = self._expand_path(output_path)
        try:
            from pydub import AudioSegment
        except ImportError:
            return "ERROR: Missing required library 'pydub'. Please install it using 'pip install pydub'."
        try:
            audio = AudioSegment.from_file(input_path)
            if operation == "convert":
                out_ext = os.path.splitext(output_path)[1].lstrip('.')
                audio.export(output_path, format=out_ext)
            elif operation == "cut":
                start_ms = kwargs.get('start_ms', 0)
                end_ms = kwargs.get('end_ms', len(audio))
                audio = audio[start_ms:end_ms]
                audio.export(output_path, format=os.path.splitext(output_path)[1].lstrip('.'))
            elif operation == "volume":
                change_db = kwargs.get('change_db', 0)
                audio = audio + change_db
                audio.export(output_path, format=os.path.splitext(output_path)[1].lstrip('.'))
            else:
                return f"ERROR: 不支持的操作 '{operation}'。"
            return f"SUCCESS: 音频处理完成，保存到 {output_path}"
        except Exception as e:
            return f"ERROR: 音频处理失败: {str(e)}"

    def video_edit(self, input_path: str, output_path: str = None, operation: str = "cut", **kwargs):
        """视频编辑"""
        input_path = self._expand_path(input_path)
        if not output_path:
            base, ext = os.path.splitext(input_path)
            output_path = f"{base}_edited{ext}"
        output_path = self._expand_path(output_path)
        try:
            try:
                from moviepy import VideoFileClip, concatenate_videoclips
                _MOVIEPY_V2 = True
            except ImportError:
                from moviepy.editor import VideoFileClip, concatenate_videoclips
                _MOVIEPY_V2 = False
        except ImportError:
            return "ERROR: Missing required library 'moviepy'. Please install it using 'pip install moviepy'."
        try:
            if operation == "cut":
                start_time = kwargs.get('start_time', 0)
                end_time = kwargs.get('end_time', None)
                clip = VideoFileClip(input_path)
                clip = clip.subclipped(start_time, end_time) if _MOVIEPY_V2 else clip.subclip(start_time, end_time)
                clip.write_videofile(output_path)
            elif operation == "concat":
                video_list = kwargs.get('videos', [])
                if not video_list:
                    return "ERROR: concat操作需要提供 videos 列表。"
                clips = [VideoFileClip(self._expand_path(v)) for v in video_list]
                final_clip = concatenate_videoclips(clips)
                final_clip.write_videofile(output_path)
            elif operation == "extract_audio":
                clip = VideoFileClip(input_path)
                clip.audio.write_audiofile(output_path)
            else:
                return f"ERROR: 不支持的操作 '{operation}'。"
            return f"SUCCESS: 视频处理完成，保存到 {output_path}"
        except Exception as e:
            return f"ERROR: 视频编辑失败: {str(e)}"

    # ==================== AI 视频生成 ====================
    def generate_video(
        self,
        prompt: str = None,
        scenes: list = None,
        template: str = None,
        output_path: str = None,
        bg_music_path: str = None,
        music_volume: float = 0.3,
        tts_text: str = None,
        tts_voice: str = None,
        images: dict = None,
        duration: int = 30,
        fps: int = 24,
        width: int = 1920,
        height: int = 1080,
    ):
        """
        AI 视频生成（V2 架构：主Agent给prompt → 子Agent设计HTML → Playwright渲染 → 合成视频）

        推荐用法：只需提供 prompt + images，工具会自动启动子Agent完成所有设计和渲染。
            generate_video(prompt="制作产品宣传片", images={"logo": "./logo.png"}, bg_music_path="./bgm.mp3")

        prompt: 视频需求描述
        images: 本地图片 {"名称": "路径", ...}，子Agent可在视频中插入这些图片
        bg_music_path: 背景音乐文件路径
        music_volume: 背景音乐音量 0.0-1.0
        tts_voice: TTS 音色
        duration: 视频总时长秒数（默认30）
        fps: 帧率（默认24）
        width/height: 视频分辨率（默认1920x1080）
        """
        try:
            from video_generator import generate_video as _gen_video
        except ImportError as e:
            return f"ERROR: 视频生成模块导入失败: {str(e)}。"

        try:
            # 展开图片路径
            expanded_images = None
            if images:
                expanded_images = {}
                for name, path in images.items():
                    expanded_images[name] = self._expand_path(path)

            if bg_music_path:
                bg_music_path = self._expand_path(bg_music_path)

            result = _gen_video(
                prompt=prompt,
                scenes=scenes,
                output_path=output_path,
                bg_music_path=bg_music_path,
                music_volume=music_volume,
                tts_voice=tts_voice,
                images=expanded_images,
                duration=duration,
                fps=fps,
                width=width,
                height=height,
            )

            if result.startswith("ERROR"):
                return result
            return f"SUCCESS: 视频生成完成，保存到 {result}"

        except Exception as e:
            return f"ERROR: 视频生成失败: {str(e)}"

    # ==================== AI 作曲 (子Agent架构) ====================
    def compose_music(
        self,
        prompt: str,
        duration: float = 60,
        key: str = "C",
        bpm: int = None,
    ):
        """
        AI 作曲（子Agent架构）：主Agent只需提供 prompt，子Agent(LLM+乐理)
        自动设计音乐结构、编写旋律、配器，最后渲染为 MP3 + MIDI。

        架构: 主Agent→子Agent(LLM+乐理)→JSON→MIDI→多乐器合成→MP3

        支持的乐器: 钢琴、小提琴、大提琴、吉他、贝斯、长笛、双簧管、
                   小号、萨克斯、弦乐合奏、合成器、合唱、架子鼓等 20+ 种

        参数:
          prompt: 音乐创作需求描述（风格、乐器、情绪、参考等）
          duration: 时长(秒)，默认60
          key: 调性 C/D/E/F/G/A/B
          bpm: 速度(可选)

        返回: MP3 文件路径（同时生成同名 .mid 文件）
        """
        try:
            from music_composer import compose_music as _compose
        except ImportError:
            return "ERROR: AI作曲模块不可用。"

        try:
            result = _compose(
                prompt=prompt,
                duration=duration,
                key=key.upper() if key else "C",
                bpm=bpm,
            )
            if result.startswith("ERROR"):
                return result
            return f"SUCCESS: AI作曲完成，MP3文件路径: {result}（同时生成MIDI文件）"
        except Exception as e:
            return f"ERROR: AI作曲失败: {str(e)}"

    # ==================== 串口通信 ====================
    def list_serial_ports(self):
        """列出可用串口"""
        try:
            import serial.tools.list_ports
        except ImportError:
            return "ERROR: Missing required library 'pyserial'. Please install it using 'pip install pyserial'."
        try:
            ports = serial.tools.list_ports.comports()
            if not ports:
                return "INFO: 未找到任何串口。"
            port_list = "\n".join([f"{port.device} - {port.description}" for port in ports])
            return f"SUCCESS: 可用串口:\n{port_list}"
        except Exception as e:
            return f"ERROR: 获取串口列表失败: {str(e)}"

    def serial_send(self, port: str, data: str, baudrate: int = 9600, timeout: float = 1.0):
        """向串口发送数据"""
        try:
            import serial
        except ImportError:
            return "ERROR: Missing required library 'pyserial'."
        try:
            with serial.Serial(port, baudrate, timeout=timeout) as ser:
                bytes_sent = ser.write(data.encode('ascii'))
            return f"SUCCESS: 已向 {port} 发送 {bytes_sent} 字节数据: {repr(data)}"
        except Exception as e:
            return f"ERROR: 发送失败: {str(e)}"

    def serial_read(self, port: str, baudrate: int = 9600, timeout: float = 1.0):
        """从串口读取一行数据"""
        try:
            import serial
        except ImportError:
            return "ERROR: Missing required library 'pyserial'."
        try:
            with serial.Serial(port, baudrate, timeout=timeout) as ser:
                line = ser.readline().decode('ascii').strip()
            if line:
                return f"SUCCESS: 从 {port} 读取到: {line}"
            else:
                return f"INFO: 从 {port} 未读取到数据（超时）"
        except Exception as e:
            return f"ERROR: 读取失败: {str(e)}"

    def delay(self, seconds: float):
        """等待指定秒数"""
        try:
            time.sleep(seconds)
            return f"SUCCESS: 已等待 {seconds} 秒。"
        except Exception as e:
            return f"ERROR: 延时失败: {str(e)}"

    # ==================== 记忆管理 ====================
    def add_long_term(self, text: str):
        """向长期记忆添加内容"""
        self.memory.add_long_term(text)
        return f"SUCCESS: 已添加到长期记忆: {text}"

    # ==================== QQ 文件发送 ====================
    def send_qq_file(self, target_type: str, target_id: str, file_path: str, message: str = ""):
        """通过 NapCat HTTP API 发送文件到 QQ"""
        file_path = self._expand_path(file_path)
        if not os.path.isfile(file_path):
            return f"ERROR: 文件不存在: {file_path}"
        if not config.napcat_http_url:
            return "ERROR: NapCat HTTP 地址未配置"
        if target_type not in ('private', 'group'):
            return f"ERROR: target_type 必须是 'private' 或 'group'"
        api_path = '/send_private_msg' if target_type == 'private' else '/send_group_msg'
        url = config.napcat_http_url.rstrip('/') + api_path

        message_segments = []
        if message:
            message_segments.append({"type": "text", "data": {"text": message}})
        ext = os.path.splitext(file_path)[1].lower()
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            message_segments.append({"type": "image", "data": {"file": file_path}})
        elif ext in ['.mp4', '.avi', '.mov', '.mkv']:
            message_segments.append({"type": "video", "data": {"file": file_path}})
        elif ext in ['.mp3', '.wav', '.flac']:
            message_segments.append({"type": "record", "data": {"file": file_path}})
        else:
            message_segments.append({"type": "file", "data": {"file": file_path}})

        params = {'message': message_segments}
        if target_type == 'private':
            params['user_id'] = int(target_id)
        else:
            params['group_id'] = int(target_id)

        headers = {}
        if config.napcat_access_token:
            headers['Authorization'] = f'Bearer {config.napcat_access_token}'
        try:
            import requests
            response = requests.post(url, json=params, headers=headers, timeout=10)
            response.raise_for_status()
            result = response.json()
            if result.get('status') == 'ok':
                return f"SUCCESS: 文件已发送至 {target_type} {target_id}"
            else:
                return f"ERROR: 发送失败: {result}"
        except Exception as e:
            return f"ERROR: 发送文件异常: {str(e)}"

    # ==================== 技能执行 ====================
    def execute_skill(self, skill_name: str, action: str = "run_script", params: dict = None):
        """执行技能脚本或获取说明"""
        if params is None:
            params = {}
        skill = self.skill_manager.get_skill(skill_name)
        if not skill:
            return f"ERROR: 技能 '{skill_name}' 不存在"
        try:
            if action == "run_script":
                script_name = params.get("script")
                if not script_name:
                    return "ERROR: 执行脚本需要提供 'script' 参数"
                result = skill.execute_script(script_name, params.get("args", []))
                return f"SUCCESS: 技能脚本执行结果:\n{result}"
            elif action == "get_info":
                content = skill.load_full_content()
                return f"SUCCESS: 技能 '{skill_name}' 的完整说明:\n{content}"
            else:
                return f"ERROR: 不支持的操作 '{action}'"
        except Exception as e:
            return f"ERROR: 执行技能异常: {str(e)}"

    # ==================== 定时任务 ====================
    def add_scheduled_task(self, message: str, trigger_type: str, trigger_args: dict):
        """添加定时任务"""
        if self.task_scheduler is None:
            return "ERROR: 定时任务模块未初始化"
        try:
            task_id = self.task_scheduler.add_task({
                'message': message,
                'trigger': trigger_type,
                'trigger_args': trigger_args,
                'enabled': True
            })
            return f"SUCCESS: 定时任务已添加，ID: {task_id}"
        except Exception as e:
            return f"ERROR: 添加任务失败: {str(e)}"

    def list_scheduled_tasks(self):
        """列出所有定时任务"""
        if self.task_scheduler is None:
            return "ERROR: 定时任务模块未初始化"
        tasks = self.task_scheduler.get_tasks()
        if not tasks:
            return "INFO: 当前没有定时任务"
        lines = []
        for t in tasks:
            lines.append(f"ID: {t['id']} | 消息: {t['message'][:30]} | 触发器: {t['trigger']} | 状态: {'启用' if t.get('enabled',True) else '禁用'}")
        return "SUCCESS: 定时任务列表:\n" + "\n".join(lines)

    def delete_scheduled_task(self, task_id: str):
        """删除定时任务"""
        if self.task_scheduler is None:
            return "ERROR: 定时任务模块未初始化"
        try:
            self.task_scheduler.remove_task(task_id)
            return f"SUCCESS: 任务 {task_id} 已删除"
        except Exception as e:
            return f"ERROR: 删除失败: {str(e)}"

    # ==================== 嵌入式开发 ====================
    def install_arduino_cli(self):
        """安装 arduino-cli"""
        import platform as pf
        import urllib.request
        import zipfile, tarfile
        try:
            subprocess.run(['arduino-cli', 'version'], capture_output=True, text=True, check=True)
            return "INFO: arduino-cli 已安装"
        except:
            pass
        system = pf.system()
        machine = pf.machine()
        base_url = "https://github.com/arduino/arduino-cli/releases/latest/download/"
        if system == "Windows":
            filename = "arduino-cli_latest_Windows_64bit.zip"
        elif system == "Linux":
            filename = "arduino-cli_latest_Linux_ARM64.tar.gz" if "arm" in machine else "arduino-cli_latest_Linux_64bit.tar.gz"
        elif system == "Darwin":
            filename = "arduino-cli_latest_macOS_64bit.tar.gz"
        else:
            return "ERROR: 不支持的操作系统"
        url = base_url + filename
        try:
            urllib.request.urlretrieve(url, filename)
        except Exception as e:
            return f"ERROR: 下载失败: {e}。请手动安装。"
        try:
            if filename.endswith('.zip'):
                with zipfile.ZipFile(filename, 'r') as zf:
                    zf.extractall('.')
            else:
                with tarfile.open(filename, 'r:gz') as tf:
                    tf.extractall('.')
            if system == "Windows":
                for f in os.listdir('.'):
                    if f.endswith('.exe') and 'arduino' in f:
                        os.rename(f, 'arduino-cli.exe')
            else:
                if os.path.exists('arduino-cli'):
                    os.chmod('arduino-cli', 0o755)
                else:
                    for f in os.listdir('.'):
                        if f.startswith('arduino-cli') and not f.endswith('.gz'):
                            os.rename(f, 'arduino-cli')
                            os.chmod('arduino-cli', 0o755)
                            break
            os.remove(filename)
            return "SUCCESS: arduino-cli 已安装到当前目录，请确保该目录在系统 PATH 中"
        except Exception as e:
            return f"ERROR: 安装失败: {str(e)}"

    def generate_ino_code(self, description: str, board_type: str):
        """生成 Arduino 代码"""
        prompt = f"""请根据以下需求生成完整的 Arduino 代码（.ino 格式），确保代码可以直接编译并上传到 {board_type}。

需求: {description}

要求：
1. 包含必要的 setup() 和 loop() 函数。
2. 适当添加注释。
3. 只输出代码，不要包含任何额外解释。
"""
        try:
            response = self.client.chat.completions.create(
                model=config.ai_model,
                messages=[
                    {"role": "system", "content": "你是 Arduino 编程专家，擅长为各种单片机生成代码。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=1.0,
                max_tokens=2000,
                **_deepseek_sdk_params(config.ai_model),
            )
            code = response.choices[0].message.content.strip()
            if code.startswith("```") and code.endswith("```"):
                lines = code.split('\n')
                if lines[0].startswith("```"): lines = lines[1:]
                if lines[-1].startswith("```"): lines = lines[:-1]
                code = '\n'.join(lines).strip()
            import tempfile
            project_name = f"esp32_rgb_{int(time.time())}"
            project_dir = os.path.join(tempfile.gettempdir(), project_name)
            os.makedirs(project_dir, exist_ok=True)
            filepath = os.path.join(project_dir, f"{project_name}.ino")
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(code)
            return f"SUCCESS: 代码已生成，项目目录: {project_dir}\n主文件: {filepath}"
        except Exception as e:
            return f"ERROR: 生成代码失败: {str(e)}"

    def compile_ino(self, project_dir: str = None, ino_path: str = None, board_fqbn: str = None, **kwargs):
        """编译 .ino 项目"""
        if project_dir is None and ino_path is not None:
            project_dir = ino_path
        if board_fqbn is None:
            return "ERROR: 缺少 board_fqbn 参数"
        if not os.path.isdir(project_dir):
            if os.path.isfile(project_dir):
                project_dir = os.path.dirname(project_dir)
            else:
                return f"ERROR: 项目目录不存在: {project_dir}"
        ino_file = None
        for f in os.listdir(project_dir):
            if f.endswith('.ino'):
                ino_file = os.path.join(project_dir, f)
                break
        if not ino_file:
            return f"ERROR: 项目目录中没有找到 .ino 文件"
        build_dir = os.path.join(project_dir, "build")
        os.makedirs(build_dir, exist_ok=True)
        result = self._run_arduino_cli(['compile', '--fqbn', board_fqbn, '--build-path', build_dir, project_dir], timeout=999)
        if result.startswith("ERROR"):
            if "platform not found" in result or "not installed" in result:
                core_name = board_fqbn.split(':')[0] + ':' + board_fqbn.split(':')[1]
                install_res = self._run_arduino_cli(['core', 'install', core_name], timeout=300)
                if install_res.startswith("ERROR"):
                    return f"ERROR: 自动安装核心失败。{install_res}\n编译错误：{result}"
                result2 = self._run_arduino_cli(['compile', '--fqbn', board_fqbn, '--build-path', build_dir, project_dir], timeout=120)
                if result2.startswith("ERROR"):
                    return f"ERROR: 编译失败，请手动上传代码。{result2}"
                return f"SUCCESS: 编译成功，构建文件位于 {build_dir}"
            return f"ERROR: 编译失败。{result}"
        return f"SUCCESS: 编译成功，构建文件位于 {build_dir}"

    def auto_detect_board_and_port(self):
        """自动检测开发板"""
        detect_res = self._run_arduino_cli(['board', 'list'])
        if detect_res.startswith("ERROR"):
            return detect_res
        lines = detect_res.strip().split('\n')
        boards = []
        for line in lines[1:]:
            parts = line.split()
            if len(parts) >= 4:
                boards.append({'port': parts[0], 'fqbn': parts[-1]})
        if not boards:
            return "INFO: 未检测到已连接的开发板。"
        return f"SUCCESS: 检测到开发板:\n" + "\n".join([f"{b['port']} -> {b['fqbn']}" for b in boards])

    def install_board_core(self, board_fqbn: str):
        """安装板卡核心"""
        update_res = self._run_arduino_cli(['core', 'update-index'])
        if update_res.startswith("ERROR"):
            return update_res
        install_res = self._run_arduino_cli(['core', 'install', board_fqbn], timeout=600)
        if install_res.startswith("ERROR"):
            return install_res
        return f"SUCCESS: 核心 {board_fqbn} 安装成功。"

    def install_arduino_library(self, library_name: str):
        """安装 Arduino 库"""
        update_res = self._run_arduino_cli(['lib', 'update-index'])
        if update_res.startswith("ERROR"):
            return update_res
        install_res = self._run_arduino_cli(['lib', 'install', library_name], timeout=120)
        if install_res.startswith("ERROR"):
            return install_res
        return f"SUCCESS: Arduino 库 {library_name} 安装成功。"

    def _run_arduino_cli(self, args, timeout=60):
        """内部执行 arduino-cli 命令"""
        try:
            result = subprocess.run(['arduino-cli'] + args, capture_output=True, text=True,
                                    timeout=timeout, encoding='utf-8', errors='replace')
            if result.returncode == 0:
                return result.stdout
            else:
                return f"ERROR: {result.stderr}"
        except FileNotFoundError:
            return "ERROR: arduino-cli 未安装"
        except subprocess.TimeoutExpired:
            return "ERROR: 命令执行超时"
        except Exception as e:
            return f"ERROR: {str(e)}"

    def save_ino_to_desktop(self, project_dir: str, filename: str = "ESP32_LED.ino"):
        """将 Arduino 项目保存到桌面"""
        if not os.path.isdir(project_dir):
            return f"ERROR: 项目目录不存在: {project_dir}"
        desktop = os.path.expanduser("~/Desktop")
        target_dir = os.path.join(desktop, os.path.basename(project_dir))
        try:
            shutil.copytree(project_dir, target_dir, dirs_exist_ok=True)
            return f"SUCCESS: 项目已保存到桌面，文件夹名: {os.path.basename(project_dir)}"
        except Exception as e:
            return f"ERROR: 保存失败: {str(e)}"

    def upload_ino(self, project_dir: str, board_fqbn: str, port: str):
        """上传已编译的项目到开发板"""
        if not os.path.isdir(project_dir):
            return f"ERROR: 项目目录不存在: {project_dir}"
        build_dir = os.path.join(project_dir, "build")
        if not os.path.isdir(build_dir):
            return "ERROR: 构建目录不存在，请先编译。"
        result = self._run_arduino_cli(['upload', '--fqbn', board_fqbn, '--port', port, '--input-dir', build_dir], timeout=60)
        if result.startswith("ERROR"):
            return result
        return f"SUCCESS: 上传成功到串口 {port}"

    # ==================== QQ 扩展工具 ====================
    def recall_message(self, message_id):
        """撤回消息"""
        if message_id == "last":
            if not hasattr(self, 'gui') or self.gui.last_sent_message_id is None:
                return "ERROR: 没有找到最近发送的消息ID"
            message_id = self.gui.last_sent_message_id
        try:
            url = config.napcat_http_url.rstrip('/') + "/delete_msg"
            payload = {"message_id": int(message_id)}
            headers = {}
            if config.napcat_access_token:
                headers['Authorization'] = f'Bearer {config.napcat_access_token}'
            response = requests.post(url, json=payload, headers=headers, timeout=10)
            response.raise_for_status()
            return "SUCCESS: 消息已撤回"
        except Exception as e:
            return f"ERROR: 撤回失败: {str(e)}"

    def add_friend(self, user_id: str, message: str = ""):
        """发送好友申请"""
        try:
            url = config.napcat_http_url.rstrip('/') + "/add_friend"
            payload = {"user_id": int(user_id), "message": message}
            headers = {}
            if config.napcat_access_token:
                headers['Authorization'] = f'Bearer {config.napcat_access_token}'
            response = requests.post(url, json=payload, headers=headers, timeout=10)
            response.raise_for_status()
            return f"SUCCESS: 好友申请已发送给 {user_id}"
        except Exception as e:
            return f"ERROR: 发送好友申请失败: {str(e)}"

    def join_group(self, group_id: str, message: str = ""):
        """申请加入群聊"""
        try:
            url = config.napcat_http_url.rstrip('/') + "/set_group_add_request"
            payload = {"group_id": int(group_id), "message": message, "type": "add"}
            headers = {}
            if config.napcat_access_token:
                headers['Authorization'] = f'Bearer {config.napcat_access_token}'
            response = requests.post(url, json=payload, headers=headers, timeout=10)
            response.raise_for_status()
            return f"SUCCESS: 已申请加入群 {group_id}"
        except Exception as e:
            return f"ERROR: 申请加群失败: {str(e)}"

    # ==================== 浏览器自动化增强（核心升级） ====================
    def _get_browser(self, proxy=None, user_agent=None):
        """
        获取当前线程的 Playwright 浏览器实例（自动创建，支持代理和自定义 UA）
        每个线程独立管理实例，解决多线程冲突
        """
        import threading
        thread_id = threading.get_ident()
        if not hasattr(self, '_browser_by_thread'):
            self._browser_by_thread = {}

        # 若已存在实例但代理参数不匹配则销毁重建
        if thread_id in self._browser_by_thread:
            browser, playwright, saved_proxy = self._browser_by_thread[thread_id]
            if proxy != saved_proxy:
                try:
                    browser.close()
                except:
                    pass
                try:
                    playwright.stop()
                except:
                    pass
                del self._browser_by_thread[thread_id]

        if thread_id not in self._browser_by_thread:
            try:
                from playwright.sync_api import sync_playwright
            except ImportError:
                return None, "ERROR: Playwright 未安装，请使用 install_python_package 安装 'playwright' 并手动运行 playwright install chromium"

            try:
                playwright = sync_playwright().start()
                launch_args = [
                    '--no-sandbox',
                    '--disable-setuid-sandbox',
                    '--disable-blink-features=AutomationControlled',
                    '--disable-features=IsolateOrigins,site-per-process',
                ]
                headful = getattr(config, 'browser_headful', False)
                browser = playwright.chromium.launch(
                    headless=not headful,
                    args=launch_args
                )
                self._browser_by_thread[thread_id] = (browser, playwright, proxy)
            except Exception as e:
                return None, f"ERROR: 启动浏览器失败: {str(e)}"

        return self._browser_by_thread[thread_id][0], None

    def _create_browser_context(self, proxy=None, user_agent=None):
        """创建带反爬参数的浏览器上下文，自动注入 webdriver 隐藏脚本"""
        browser, err = self._get_browser(proxy=proxy)
        if err:
            return None, err

        if user_agent is None:
            ua_list = [
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
                "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
            ]
            user_agent = random.choice(ua_list)

        context_options = {
            "user_agent": user_agent,
            "locale": "en-US",
            "timezone_id": "America/New_York",
        }
        if proxy:
            context_options["proxy"] = {"server": proxy}

        context = browser.new_context(**context_options)
        # 去除 navigator.webdriver 标记
        context.add_init_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        return context, None

    def browser_navigate(self, url, wait_until="networkidle", proxy=None, user_agent=None, retries=2):
        """
        打开浏览器并导航到指定 URL，支持代理、自定义 UA、失败重试
        """
        for attempt in range(retries + 1):
            try:
                context, err = self._create_browser_context(proxy=proxy, user_agent=user_agent)
                if err:
                    return err
                page = context.new_page()
                self._current_page = page
                page.goto(url, wait_until=wait_until, timeout=30000)
                return f"SUCCESS: 已打开 {url}"
            except Exception as e:
                error_str = str(e)
                if attempt < retries and ("timeout" in error_str or "net::" in error_str or "NS_ERROR" in error_str):
                    time.sleep(2)
                    continue
                if "cannot switch to a different thread" in error_str:
                    self.browser_close()
                    continue
                return f"ERROR: 导航失败: {error_str}"

    def browser_close(self):
        """关闭当前线程的浏览器实例并清理资源"""
        import threading
        thread_id = threading.get_ident()
        if hasattr(self, '_browser_by_thread') and thread_id in self._browser_by_thread:
            browser, playwright, _ = self._browser_by_thread.pop(thread_id)
            try:
                browser.close()
            except:
                pass
            try:
                playwright.stop()
            except:
                pass
        if hasattr(self, '_current_page'):
            del self._current_page
        return "SUCCESS: 浏览器已关闭"

    def browser_set_proxy(self, proxy_server: str):
        """
        为当前线程设置代理（如 'http://127.0.0.1:7890' 或 'socks5://127.0.0.1:10808'）
        设置后下一次 browser_navigate 将自动使用该代理
        """
        import threading
        thread_id = threading.get_ident()
        if not hasattr(self, '_browser_proxies'):
            self._browser_proxies = {}
        self._browser_proxies[thread_id] = proxy_server
        return f"SUCCESS: 代理已设置为 {proxy_server}，下次导航生效"

    def _get_thread_proxy(self):
        """获取当前线程配置的代理"""
        import threading
        return getattr(self, '_browser_proxies', {}).get(threading.get_ident())

    def browser_click(self, selector, min_delay=0.5, max_delay=2.0):
        """模拟人类点击：随机移动鼠标并延时，降低反爬检测"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        page = self._current_page
        try:
            time.sleep(random.uniform(min_delay, max_delay))
            element = page.wait_for_selector(selector, timeout=5000)
            box = element.bounding_box()
            if box:
                page.mouse.move(box['x'] + box['width']/2, box['y'] + box['height']/2)
            element.click()
            return f"SUCCESS: 已点击 {selector}"
        except Exception as e:
            return f"ERROR: 点击失败: {str(e)}"

    def browser_fill(self, selector, text, clear_first=True, delay=0.15):
        """模拟人类逐字输入，可设置每字间隔（秒），防止触发机器人检测"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        page = self._current_page
        try:
            if clear_first:
                page.fill(selector, '')
            for char in text:
                page.type(selector, char, delay=random.uniform(50, 150))  # 毫秒
            return f"SUCCESS: 已填写 {text} 到 {selector}"
        except Exception as e:
            return f"ERROR: 填写失败: {str(e)}"

    def browser_get_text(self, selector, timeout=5000):
        """获取匹配选择器的元素文本内容，支持等待超时"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            element = self._current_page.wait_for_selector(selector, timeout=timeout)
            text = element.text_content()
            return f"SUCCESS: 文本内容: {text.strip() if text else '（空）'}"
        except Exception as e:
            return f"ERROR: 获取文本失败: {str(e)}"

    def browser_evaluate(self, script, safe=True):
        """在页面中执行 JavaScript 脚本，可获取动态数据（安全模式下限制危险操作）"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        if safe and not self._is_safe_js(script):
            return "ERROR: 安全模式拒绝执行可能危险的脚本"
        try:
            result = self._current_page.evaluate(script)
            return f"SUCCESS: 执行结果: {result}"
        except Exception as e:
            return f"ERROR: 脚本执行失败: {str(e)}"

    # AI 常用别名（这些不是独立工具，只是 browser_evaluate 的别名）
    def browser_execute_js(self, script, safe=True):
        return self.browser_evaluate(script, safe)

    def execute_js(self, script, safe=True):
        return self.browser_evaluate(script, safe)

    def browser_get_element_html(self, selector):
        """获取指定元素的 outerHTML 内容"""
        return self.browser_evaluate(
            f'(()=>{{const el=document.querySelector({selector!r});return el?el.outerHTML:null}})()',
            safe=True
        )

    def _is_safe_js(self, script: str) -> bool:
        """简单的 JavaScript 安全检查（仅允许只读操作）"""
        dangerous_keywords = [
            'document.write', 'window.open', 'alert', 'confirm', 'prompt',
            'localStorage', 'sessionStorage', 'indexedDB', 'XMLHttpRequest',
            'fetch', 'WebSocket', 'postMessage', 'eval', 'Function',
            'setTimeout', 'setInterval', 'location.href', 'location.replace',
            'document.cookie', 'innerHTML', 'outerHTML', 'insertAdjacentHTML',
            'createElement', 'appendChild', 'removeChild', 'replaceChild'
        ]
        script_lower = script.lower()
        for kw in dangerous_keywords:
            if kw in script_lower:
                return False
        return True

    def browser_screenshot(self, path: str = "./screenshots/browser.png"):
        """截取当前页面视图（非全页）"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            self._current_page.screenshot(path=path)
            return f"SUCCESS: 截图已保存到 {path}"
        except Exception as e:
            return f"ERROR: 截图失败: {str(e)}"

    # -------------------- 高级网页操作 --------------------
    def browser_submit_form(self, form_selector=None):
        """提交当前焦点所在表单或指定表单并等待页面加载完成"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        page = self._current_page
        try:
            if form_selector:
                page.locator(form_selector).evaluate("el => el.submit()")
            else:
                page.keyboard.press('Enter')
            page.wait_for_load_state('networkidle')
            return "SUCCESS: 表单已提交"
        except Exception as e:
            return f"ERROR: 提交失败: {str(e)}"

    def browser_wait_for_navigation(self, timeout=10000):
        """等待页面导航完成（网络空闲），用于点击链接后等待加载"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            self._current_page.wait_for_load_state('networkidle', timeout=timeout)
            return "SUCCESS: 导航完成"
        except Exception as e:
            return f"ERROR: {str(e)}"

    def browser_upload_file(self, selector, file_path):
        """
        上传文件到指定 input[type=file] 元素，可用于发视频、图片等
        """
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        if not os.path.isabs(file_path):
            file_path = os.path.abspath(file_path)
        if not os.path.isfile(file_path):
            return f"ERROR: 文件不存在: {file_path}"
        try:
            page = self._current_page
            file_input = page.wait_for_selector(selector, timeout=5000)
            file_input.set_input_files(file_path)
            return f"SUCCESS: 已上传 {file_path}"
        except Exception as e:
            return f"ERROR: 上传失败: {str(e)}"

    def browser_extract_all_links(self, base_selector="a"):
        """
        提取页面中所有链接（a 标签的 href），返回列表
        """
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            page = self._current_page
            links = page.eval_on_selector_all(
                base_selector,
                "elements => elements.map(el => el.href).filter(href => href)"
            )
            return f"SUCCESS: 提取到 {len(links)} 个链接:\n" + "\n".join(links[:50]) + ("\n..." if len(links) > 50 else "")
        except Exception as e:
            return f"ERROR: {str(e)}"

    def browser_screenshot_full_page(self, path="./screenshots/fullpage.png"):
        """
        截取整个页面（包括滚动区域），保存为 PNG
        """
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            self._current_page.screenshot(path=path, full_page=True)
            return f"SUCCESS: 全页截图已保存到 {path}"
        except Exception as e:
            return f"ERROR: 截图失败: {str(e)}"

    def browser_wait_for_selector(self, selector, timeout=10000, state="visible"):
        """
        等待指定选择器的元素出现（默认等待可见），超时返回错误
        """
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            self._current_page.wait_for_selector(selector, timeout=timeout, state=state)
            return f"SUCCESS: 元素 {selector} 已出现"
        except Exception as e:
            return f"ERROR: 等待元素超时: {str(e)}"

    def browser_download_images(self, selector="img", output_dir="./downloads/images", attr="src"):
        """
        下载当前页面中所有匹配选择器的图片，保存到指定目录
        """
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        page = self._current_page
        os.makedirs(output_dir, exist_ok=True)
        downloaded = []
        try:
            elements = page.query_selector_all(selector)
            for i, el in enumerate(elements):
                src = el.get_attribute(attr)
                if not src:
                    continue
                if src.startswith("//"):
                    src = "https:" + src
                elif src.startswith("/"):
                    from urllib.parse import urljoin
                    src = urljoin(page.url, src)
                try:
                    r = requests.get(src, timeout=15)
                    if r.status_code == 200:
                        ext = os.path.splitext(src.split("?")[0])[1] or ".png"
                        filename = f"img_{i+1}{ext}"
                        filepath = os.path.join(output_dir, filename)
                        with open(filepath, 'wb') as f:
                            f.write(r.content)
                        downloaded.append(filepath)
                except:
                    continue
            return f"SUCCESS: 已下载 {len(downloaded)} 张图片到 {output_dir}: " + "; ".join(downloaded)
        except Exception as e:
            return f"ERROR: 下载图片失败: {str(e)}"

    def browser_switch_to_iframe(self, selector):
        """
        将操作焦点切换到指定的 iframe 内部，以便操作其中的元素
        """
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            frame = self._current_page.wait_for_selector(selector, timeout=5000)
            self._current_page = frame.content_frame()
            return f"SUCCESS: 已切换到 iframe: {selector}"
        except Exception as e:
            return f"ERROR: 切换 iframe 失败: {str(e)}"

    def browser_switch_to_main_frame(self):
        """
        返回主框架（从 iframe 切出）
        """
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            self._current_page = self._current_page.page.main_frame
            return "SUCCESS: 已返回主框架"
        except Exception as e:
            return f"ERROR: 返回主框架失败: {str(e)}"

    def browser_execute_workflow(self, workflow_json):
        """
        按 JSON 描述的工作流逐步执行浏览器操作，减少 AI 调用次数
        支持 action: navigate, fill, click, wait, wait_selector, extract_text, download_images, evaluate, screenshot
        """
        if not hasattr(self, '_current_page') and workflow_json and workflow_json[0].get("action") != "navigate":
            return "ERROR: 当前没有打开页面，工作流第一步必须是 navigate"
        results = []
        for step in workflow_json:
            action = step.get("action")
            try:
                if action == "navigate":
                    res = self.browser_navigate(step["url"])
                    results.append(res)
                elif action == "fill":
                    res = self.browser_fill(step["selector"], step["text"])
                    results.append(res)
                elif action == "click":
                    res = self.browser_click(step["selector"])
                    results.append(res)
                elif action == "wait":
                    time.sleep(float(step.get("seconds", 1)))
                    results.append(f"SUCCESS: 已等待 {step.get('seconds', 1)} 秒")
                elif action == "wait_selector":
                    res = self.browser_wait_for_selector(step["selector"], int(step.get("timeout", 10000)))
                    results.append(res)
                elif action == "extract_text":
                    res = self.browser_get_text(step["selector"])
                    results.append(res)
                elif action == "download_images":
                    res = self.browser_download_images(
                        step.get("selector", "img"),
                        step.get("output_dir", "./downloads/images")
                    )
                    results.append(res)
                elif action == "evaluate":
                    res = self.browser_evaluate(step["script"], step.get("safe", True))
                    results.append(res)
                elif action == "screenshot":
                    res = self.browser_screenshot(step.get("path", "./screenshots/browser.png"))
                    results.append(res)
                else:
                    results.append(f"ERROR: 未知步骤动作 {action}")
            except Exception as e:
                results.append(f"ERROR: 执行步骤 {action} 失败: {str(e)}")
        return "SUCCESS: 工作流执行完毕\n" + "\n".join(results)

    # -------------------- 网络代理与翻墙工具 --------------------
    def setup_system_proxy(self, proxy_server: str):
        """
        设置当前 Python 进程的 HTTP/HTTPS 代理（影响 requests 等库）
        浏览器代理请使用 browser_set_proxy
        """
        os.environ['HTTP_PROXY'] = proxy_server
        os.environ['HTTPS_PROXY'] = proxy_server
        return f"SUCCESS: 系统代理已设置为 {proxy_server}"

    def install_v2ray(self, config_url: str):
        """
        在 Windows 上自动下载 v2ray 核心并使用指定配置文件启动
        新增：当 GitHub 下载失败时，尝试从国内镜像下载
        """
        if sys.platform != 'win32':
            return "INFO: 当前仅支持 Windows 自动化安装，其他系统请手动配置"
        import tempfile, urllib.request, zipfile
        
        v2ray_dir = os.path.expanduser("~/v2ray-core")
        os.makedirs(v2ray_dir, exist_ok=True)
        zip_path = os.path.join(v2ray_dir, "v2ray.zip")

        # 下载源列表（按优先级）
        download_urls = [
            # 官方源（大概率失败）
            "https://github.com/v2fly/v2ray-core/releases/download/v5.1.0/v2ray-windows-64.zip",
            # 尝试使用 ghproxy 加速
            "https://ghproxy.com/https://github.com/v2fly/v2ray-core/releases/download/v5.1.0/v2ray-windows-64.zip",
            # 尝试 mirror.ghproxy.com
            "https://mirror.ghproxy.com/https://github.com/v2fly/v2ray-core/releases/download/v5.1.0/v2ray-windows-64.zip",
        ]

        download_success = False
        for url in download_urls:
            try:
                urllib.request.urlretrieve(url, zip_path)
                download_success = True
                break
            except Exception as e:
                continue

        if not download_success:
            return ("ERROR: v2ray 核心下载失败。所有下载源均无法连接。\n"
                    "建议手动下载 v2ray-windows-64.zip 并解压到 {} 目录，然后重新运行。".format(v2ray_dir))

        # 解压
        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(v2ray_dir)
            # 删除压缩包
            os.remove(zip_path)
        except Exception as e:
            return f"ERROR: 解压失败: {e}"

        # 下载配置文件（同样尝试代理下载）
        config_file = os.path.join(v2ray_dir, "config.json")
        try:
            urllib.request.urlretrieve(config_url, config_file)
        except Exception as e:
            return (f"ERROR: 配置文件下载失败: {e}\n"
                    f"v2ray 核心已安装到 {v2ray_dir}，请手动放入 config.json 文件后启动。")

        # 启动 v2ray
        try:
            subprocess.Popen(
                [os.path.join(v2ray_dir, "v2ray.exe"), "run", "-config", config_file],
                shell=False
            )
            return ("SUCCESS: v2ray 已安装并启动。\n"
                    "本地代理地址通常为 socks5://127.0.0.1:10808 或 http://127.0.0.1:10809\n"
                    "请使用 browser_set_proxy 工具设置浏览器代理。")
        except Exception as e:
            return f"ERROR: v2ray 启动失败: {e}"

    def setup_clash_agent(self, subscription_url: str):
        """
        提供 Clash Verge 安装指引，方便 AI 指导用户配置更强大的分流代理
        """
        return (
            "INFO: 自动安装 Clash GUI 较复杂，请指引用户按以下步骤操作：\n"
            "1. 从 https://github.com/clash-verge-rev/clash-verge-rev/releases 下载 Windows 版本\n"
            f"2. 安装后导入订阅链接: {subscription_url}\n"
            "3. 启动代理，默认端口 7890\n"
            "4. 使用 browser_set_proxy 工具设置代理: http://127.0.0.1:7890"
        )

    # ==================== 其他原有工具继续保留 ====================
    def generate_docx(self, filepath: str, title: str, content: str):
        """生成 Word 文档（.docx）"""
        try:
            from docx import Document
        except ImportError:
            return "ERROR: Missing required library 'python-docx'. Please install it using 'pip install python-docx'."
        try:
            doc = Document()
            doc.add_heading(title, level=1)
            for paragraph in content.split('\n'):
                doc.add_paragraph(paragraph)
            doc.save(filepath)
            return f"SUCCESS: Word 文档已保存到 {filepath}"
        except Exception as e:
            return f"ERROR: 生成文档失败: {str(e)}"

    def send_voice(self, target_type: str, target_id: str, voice_file: str, voice_format: str = "amr"):
        """发送语音消息到 QQ"""
        if not config.napcat_http_url:
            return "ERROR: NapCat HTTP 地址未配置"
        local_path = None
        if voice_file.startswith(('http://', 'https://')):
            import tempfile
            fd, local_path = tempfile.mkstemp(suffix=".mp3")
            os.close(fd)
            try:
                r = requests.get(voice_file, stream=True, timeout=30)
                r.raise_for_status()
                with open(local_path, 'wb') as f:
                    for chunk in r.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
            except Exception as e:
                return f"ERROR: 下载语音文件失败: {str(e)}"
        else:
            local_path = self._expand_path(voice_file)
            if not os.path.isfile(local_path):
                return f"ERROR: 语音文件不存在: {local_path}"

        segments = [{"type": "record", "data": {"file": local_path}}]
        if target_type == 'private':
            api_path = '/send_private_msg'
            params = {'user_id': int(target_id)}
        else:
            api_path = '/send_group_msg'
            params = {'group_id': int(target_id)}
        params['message'] = segments

        url = config.napcat_http_url.rstrip('/') + api_path
        headers = {}
        if config.napcat_access_token:
            headers['Authorization'] = f'Bearer {config.napcat_access_token}'
        try:
            response = requests.post(url, json=params, headers=headers, timeout=10)
            response.raise_for_status()
            result = response.json()
            if result.get('status') == 'ok':
                return f"SUCCESS: 语音已发送至 {target_type} {target_id}"
            else:
                return f"ERROR: 发送失败: {result}"
        except Exception as e:
            return f"ERROR: 发送异常: {str(e)}"
        finally:
            if voice_file.startswith(('http://', 'https://')) and local_path and os.path.exists(local_path):
                os.unlink(local_path)

    def text_to_voice(self, text: str, output_path: str = None, voice: str = "zh-CN-XiaoxiaoNeural"):
        """文本转语音（edge-tts）"""
        try:
            import edge_tts
        except ImportError:
            return "ERROR: 缺少 edge-tts 库，请执行 'pip install edge-tts'"
        if not output_path:
            import tempfile
            fd, output_path = tempfile.mkstemp(suffix=".mp3")
            os.close(fd)
        try:
            async def _tts():
                communicate = edge_tts.Communicate(text, voice)
                await communicate.save(output_path)
            import asyncio
            asyncio.run(_tts())
            return f"SUCCESS: 语音文件已保存至 {output_path}"
        except Exception as e:
            return f"ERROR: 语音合成失败: {str(e)}"

    def send_text_voice(self, target_type: str, target_id: str, text: str, voice: str = "zh-CN-XiaoxiaoNeural"):
        """将文本转换为语音并发送到QQ"""
        tts_result = self.text_to_voice(text, voice=voice)
        if tts_result.startswith("ERROR"):
            return tts_result
        file_path = tts_result.split("至 ")[1].strip()
        return self.send_voice(target_type, target_id, file_path)

    def send_music_as_voice(self, target_type: str, target_id: str, music_url: str):
        """从网络下载音乐并以语音形式发送到QQ"""
        return self.send_voice(target_type, target_id, music_url)

    def send_to(self, target_type: str, target_id: str, content: str):
        """向QQ发送消息（支持文本、图片、文件、语音、表情）"""
        if not config.napcat_http_url:
            return "ERROR: NapCat HTTP 地址未配置"
        segments = self._parse_message_to_segments(content)
        if target_type == 'private':
            api_path = '/send_private_msg'
            params = {'user_id': int(target_id)}
        else:
            api_path = '/send_group_msg'
            params = {'group_id': int(target_id)}
        params['message'] = segments
        url = config.napcat_http_url.rstrip('/') + api_path
        headers = {}
        if config.napcat_access_token:
            headers['Authorization'] = f'Bearer {config.napcat_access_token}'
        try:
            response = requests.post(url, json=params, headers=headers, timeout=10)
            response.raise_for_status()
            result = response.json()
            if result.get('status') == 'ok':
                return f"SUCCESS: 消息已发送至 {target_type} {target_id}"
            else:
                return f"ERROR: 发送失败: {result}"
        except Exception as e:
            return f"ERROR: 发送异常: {str(e)}"

    def _parse_message_to_segments(self, message):
        """将包含特殊标记的字符串解析为消息段数组"""
        segments = []
        pattern = r'\[(IMAGE|FILE|VOICE|FACE):([^\]]+)\]'
        last_end = 0
        for match in re.finditer(pattern, message):
            start, end = match.span()
            if start > last_end:
                text_part = message[last_end:start]
                if text_part:
                    segments.append({"type": "text", "data": {"text": text_part}})
            tag, value = match.groups()
            if tag == "IMAGE":
                abs_path = os.path.abspath(os.path.expanduser(value))
                if os.path.isfile(abs_path):
                    segments.append({"type": "image", "data": {"file": abs_path}})
                else:
                    segments.append({"type": "text", "data": {"text": f"[图片文件不存在: {value}]"}})
            elif tag == "FILE":
                abs_path = os.path.abspath(os.path.expanduser(value))
                if os.path.isfile(abs_path):
                    segments.append({"type": "file", "data": {"file": abs_path}})
                else:
                    segments.append({"type": "text", "data": {"text": f"[文件不存在: {value}]"}})
            elif tag == "VOICE":
                abs_path = os.path.abspath(os.path.expanduser(value))
                if os.path.isfile(abs_path):
                    segments.append({"type": "record", "data": {"file": abs_path}})
                else:
                    segments.append({"type": "text", "data": {"text": f"[语音文件不存在: {value}]"}})
            elif tag == "FACE":
                try:
                    face_id = int(value)
                    segments.append({"type": "face", "data": {"id": face_id}})
                except:
                    segments.append({"type": "text", "data": {"text": f"[无效表情ID: {value}]"}})
            last_end = end
        if last_end < len(message):
            remaining = message[last_end:]
            if remaining:
                segments.append({"type": "text", "data": {"text": remaining}})
        if not segments:
            segments.append({"type": "text", "data": {"text": message}})
        return segments

    # ==================== 物联网设备管理 ====================
    def query_devices(self):
        """查询所有已注册的物联网设备"""
        from IOT_manager import iot_manager
        devices = iot_manager.list_devices()
        if not devices:
            return "INFO: 当前没有注册任何物联网设备。"
        lines = []
        for idx, dev in enumerate(devices, 1):
            line = f"{idx}. {dev['name']} - 类型: {dev['type']} - 协议: {dev['protocol']}"
            if dev['type'] == 'complex' and dev.get('presets'):
                preset_names = [p['name'] for p in dev['presets']]
                line += f" - 可用指令: {', '.join(preset_names)}"
            if dev.get('notes'):
                line += f" - 注意事项: {dev['notes']}"
            lines.append(line)
        return "SUCCESS: 已注册设备列表:\n" + "\n".join(lines)

    def control_bool_device(self, device_name: str, state: str):
        """控制布尔类设备（开关）"""
        from IOT_manager import iot_manager
        if state.lower() not in ('on', 'off'):
            return f"ERROR: 无效状态 '{state}'，布尔设备只接受 on/off"
        dev = iot_manager.get_device(device_name)
        if not dev:
            return f"ERROR: 设备 '{device_name}' 不存在"
        if dev.device_type != 'bool':
            return f"ERROR: 设备 '{device_name}' 不是布尔类型，请使用 control_complex_device"
        return iot_manager.send_to_device(device_name, state)

    def control_complex_device(self, device_name: str, command: str):
        """控制复杂类设备"""
        from IOT_manager import iot_manager
        dev = iot_manager.get_device(device_name)
        if not dev:
            return f"ERROR: 设备 '{device_name}' 不存在"
        if dev.device_type != 'complex':
            return f"ERROR: 设备 '{device_name}' 不是复杂类型，请使用 control_bool_device"
        return iot_manager.send_to_device(device_name, command)

    def ai_add_trigger(self, name: str, sensor_name: str, tasks: list, match_pattern: str = ""):
        """AI 动态添加触发器"""
        from IOT_manager import iot_manager
        sensor = iot_manager.sensors.get(sensor_name)
        if not sensor:
            return f"ERROR: 传感器 '{sensor_name}' 不存在"
        normalized_tasks = []
        for task in tasks:
            ttype = task.get('type')
            if ttype == 'control_bool_device':
                normalized_tasks.append({
                    'type': 'control_device',
                    'device_name': task.get('device_name'),
                    'command': task.get('command')
                })
            else:
                normalized_tasks.append(task)
        trigger_data = {
            "name": name,
            "sensor_name": sensor_name,
            "match_pattern": match_pattern,
            "tasks": normalized_tasks,
            "enabled": True
        }
        if iot_manager.add_trigger(trigger_data):
            return f"SUCCESS: 触发器 '{name}' 已添加"
        else:
            return f"ERROR: 触发器名称 '{name}' 已存在"

    def delete_trigger(self, trigger_name: str):
        """删除触发器"""
        from IOT_manager import iot_manager
        if iot_manager.remove_trigger(trigger_name):
            return f"SUCCESS: 触发器 '{trigger_name}' 已删除"
        else:
            return f"ERROR: 触发器 '{trigger_name}' 不存在"

    def search_baidu(self, query: str, num_results: int = 5):
        """
        使用百度搜索（免 API），返回前 num_results 条结果的标题和链接。
        注意：频繁调用可能触发反爬，用于 AI 临时查找信息。
        """
        try:
            import requests
            from bs4 import BeautifulSoup
        except ImportError:
            return "ERROR: Missing required libraries 'requests' and 'beautifulsoup4'. Please install them."

        headers = {
            "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                           "AppleWebKit/537.36 (KHTML, like Gecko) "
                           "Chrome/124.0.0.0 Safari/537.36")
        }
        url = f"https://www.baidu.com/s?wd={requests.utils.quote(query)}"
        try:
            resp = requests.get(url, headers=headers, timeout=10)
            resp.raise_for_status()
        except Exception as e:
            return f"ERROR: 百度搜索请求失败: {str(e)}"

        soup = BeautifulSoup(resp.text, 'html.parser')
        results = []
        for item in soup.select('.result.c-container')[:num_results]:
            title_tag = item.select_one('h3 a')
            if not title_tag:
                continue
            title = title_tag.get_text(strip=True)
            link = title_tag.get('href', '')
            # 百度链接是经过跳转的，可以保留
            abstract_tag = item.select_one('.c-abstract')
            abstract = abstract_tag.get_text(strip=True) if abstract_tag else ''
            results.append(f"{title}\n  链接: {link}\n  摘要: {abstract}")

        if not results:
            return "INFO: 百度搜索未找到结果，请尝试更换关键词。"
        return "SUCCESS: 百度搜索结果:\n" + "\n".join(results)

    # ==================== 增强爬虫工具 ====================
    def _summarize_with_ai(self, content: str, content_type: str = "HTML", purpose: str = "") -> str:
        """
        AI 摘要中间层：将大量原始代码/文本提交给 AI 做信息提炼，
        只返回关键信息摘要，避免主 Agent 上下文被海量 HTML 撑爆。
        """
        try:
            api_key = getattr(config, 'ai_api_key', '')
            base_url = getattr(config, 'ai_base_url', '')
            model = getattr(config, 'ai_model', 'gpt-3.5-turbo')
            if not api_key or not base_url:
                return None

            prompt = f"""你是一个专业的信息提取助手。请分析以下{content_type}代码，提取并整理关键信息。

要求：
1. 提取页面中的核心文本内容（标题、正文、关键描述）
2. 提取所有重要的 URL 链接（用完整的绝对路径），分类整理为导航链接、资源链接、外部链接等
3. 识别结构化数据：表格、列表、价格信息、日期、联系方式等
4. 忽略无关的样式代码、广告脚本、埋点代码
5. 用简洁的中文分点总结，控制在 2000 字符以内
{purpose}

{content_type}代码：
{content[:60000]}"""

            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": model,
                "messages": [
                    {"role": "system", "content": "你是一个专业的信息提取助手。只返回提取到的关键信息，不闲聊。用中文回答。"},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.3,
                "max_tokens": 2000
            }
            # DeepSeek 深度思考模式
            if 'deepseek' in str(model).lower() and getattr(config, 'deepseek_thinking_enabled', False):
                payload["thinking"] = {"type": "enabled"}
                payload["reasoning_effort"] = getattr(config, 'deepseek_reasoning_effort', 'high')
                ctx = getattr(config, 'deepseek_context_window', 0)
                if ctx:
                    payload["max_tokens"] = ctx
            response = requests.post(
                f"{base_url}/chat/completions",
                headers=headers,
                json=payload,
                timeout=30
            )
            if response.status_code == 200:
                summary = response.json()["choices"][0]["message"]["content"]
                return summary
            return None
        except Exception as e:
            return None

    def browser_get_page_html(self, max_chars=50000, summarize=True):
        """获取当前页面完整 HTML 源码。summarize=True 时自动调用 AI 提取关键信息后返回摘要。"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            html = self._current_page.content()
            original_len = len(html)
            if summarize and original_len > 3000:
                summary = self._summarize_with_ai(html, "HTML", "注意保留页面中所有有价值的 URL 链接。")
                if summary:
                    return f"SUCCESS: 页面 AI 分析摘要 (原始 {original_len} 字符，已浓缩):\n{summary}\n\n💡 如需查看原始 HTML 中的特定区域，使用 browser_evaluate 定位元素。"
            if len(html) > max_chars:
                html = html[:max_chars] + f"\n\n... (已截断，原始长度: {original_len} 字符。可用 browser_evaluate 获取特定数据)"
            return f"SUCCESS: 页面 HTML ({len(html)} 字符)\n{html}"
        except Exception as e:
            return f"ERROR: {str(e)}"

    def browser_extract_all_text(self, max_chars=8000):
        """提取页面中所有可见文本内容（截断），快速了解页面信息"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            text = self._current_page.evaluate("""() => {
                const skipTags = ['SCRIPT', 'STYLE', 'NOSCRIPT', 'IFRAME', 'SVG'];
                const walker = document.createTreeWalker(document.body, NodeFilter.SHOW_TEXT);
                let result = [];
                while(walker.nextNode()) {
                    const parent = walker.currentNode.parentElement;
                    if (!parent || skipTags.includes(parent.tagName)) continue;
                    const t = walker.currentNode.textContent.trim();
                    if (t) result.push(t);
                }
                return result.join('\\n');
            }""")
            if len(text) > max_chars:
                text = text[:max_chars] + f"\n...(已截断 {max_chars}/{len(text)} 字符)"
            return f"SUCCESS: 页面文本\n{text}"
        except Exception as e:
            return f"ERROR: {str(e)}"

    def browser_scroll_and_extract(self, selector, scroll_times=5, scroll_delay=1.5, max_items=50):
        """滚动页面并多次提取数据，适合无限滚动/懒加载页面（如电商列表、社交媒体）"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        page = self._current_page
        all_items = []
        seen = set()
        try:
            for i in range(scroll_times):
                items = page.eval_on_selector_all(selector, """elements => elements.map(el => {
                    const text = el.innerText || el.textContent || '';
                    const href = el.querySelector('a') ? el.querySelector('a').href : '';
                    const img = el.querySelector('img') ? el.querySelector('img').src : '';
                    return {text: text.trim().substring(0, 300), href: href, img: img};
                })""")
                for item in items:
                    key = item.get('text', '')[:50]
                    if key and key not in seen:
                        seen.add(key)
                        all_items.append(item)
                if len(all_items) >= max_items:
                    break
                page.evaluate(f"window.scrollBy(0, {800 + i*200})")
                time.sleep(scroll_delay)
            result = f"提取到 {len(all_items)} 条数据:\n"
            for idx, item in enumerate(all_items[:max_items], 1):
                result += f"{idx}. {item['text'][:200]}\n"
                if item.get('href'):
                    result += f"   链接: {item['href']}\n"
            if len(all_items) > max_items:
                result += f"...（共 {len(all_items)} 条，仅显示前 {max_items} 条）"
            return f"SUCCESS: {result}"
        except Exception as e:
            return f"ERROR: 滚动提取失败: {str(e)}"

    def browser_extract_table(self, selector="table", max_rows=100, output_file=None):
        """提取页面中 HTML 表格数据，可选保存为 CSV"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            data = self._current_page.evaluate(f"""(sel) => {{
                const table = document.querySelector(sel);
                if (!table) return null;
                const headers = [];
                table.querySelectorAll('thead th, thead td, tr:first-child th, tr:first-child td').forEach(h => headers.push(h.innerText.trim()));
                const rows = [];
                table.querySelectorAll('tbody tr, tr').forEach((tr, idx) => {{
                    if (idx === 0 && headers.length > 0) {{
                        const firstCells = Array.from(tr.querySelectorAll('th, td')).map(c => c.innerText.trim());
                        if (JSON.stringify(firstCells) === JSON.stringify(headers)) return;
                    }}
                    const cells = Array.from(tr.querySelectorAll('td, th')).map(c => c.innerText.trim().substring(0, 200));
                    if (cells.length > 0) rows.push(cells);
                }});
                return {{headers, rows: rows.slice(0, {max_rows})}};
            }}""", selector)
            if not data:
                return f"ERROR: 未找到表格 {selector}"
            lines = []
            if data['headers']:
                lines.append(" | ".join(data['headers']))
                lines.append("-" * 40)
            for row in data['rows']:
                lines.append(" | ".join(row))
            result = "\n".join(lines)
            total_rows = len(data['rows'])
            result = f"表格数据 (共{total_rows}行):\n{result}"
            if output_file:
                import csv
                with open(output_file, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)
                    if data['headers']:
                        writer.writerow(data['headers'])
                    writer.writerows(data['rows'])
                result += f"\n已保存到 CSV: {output_file}"
            return f"SUCCESS: {result}"
        except Exception as e:
            return f"ERROR: 提取表格失败: {str(e)}"

    def browser_extract_json_from_script(self, key_pattern=None):
        """从页面 <script> 标签中提取 JSON 数据（常用于 SPA 页面）"""
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面"
        try:
            patterns = self._current_page.evaluate("""() => {
                const scripts = document.querySelectorAll('script[type="application/json"], script[id*="__NEXT"], script[id*="__NUXT"], script[type="application/ld+json"]');
                const results = [];
                scripts.forEach(s => {
                    try {
                        const data = JSON.parse(s.textContent);
                        results.push({id: s.id, type: s.type, keys: Object.keys(data).slice(0, 10)});
                    } catch(e) {}
                });
                return results;
            }""")
            if not patterns:
                return "INFO: 未找到嵌入式 JSON 数据块"
            lines = ["找到以下 JSON 数据块:"]
            for p in patterns:
                lines.append(f"  - id={p['id'] or '(无)'} type={p['type'] or '(无)'} 顶层键: {p['keys']}")
            lines.append("\n使用 browser_evaluate 提取具体数据，例如:")
            lines.append("  browser_evaluate('JSON.parse(document.querySelector(\"script#__NEXT_DATA__\").textContent)')")
            return "SUCCESS: " + "\n".join(lines)
        except Exception as e:
            return f"ERROR: {str(e)}"

    def browser_navigate_smart(self, url, wait_until="networkidle", timeout=45000):
        """智能导航：自动初始化浏览器+重试+超时降级+错误提示（推荐作为首选导航工具）"""
        # 如果浏览器未启动，自动初始化
        if not hasattr(self, '_current_page') and not hasattr(self, '_browser_by_thread'):
            result = self.browser_navigate(url, wait_until=wait_until)
            if result.startswith("SUCCESS"):
                return f"SUCCESS: 已自动初始化浏览器并导航到 {url}"
            return result
        page = getattr(self, '_current_page', None)
        if page:
            try:
                page.goto(url, wait_until=wait_until, timeout=timeout)
                return f"SUCCESS: 已导航到 {url}，当前标题: {page.title()}"
            except Exception as e:
                error_str = str(e)
                if "timeout" in error_str.lower():
                    try:
                        page.goto(url, wait_until="domcontentloaded", timeout=30000)
                        time.sleep(3)
                        return f"SUCCESS: 页面部分加载（网络较慢），URL: {url}，标题: {page.title()}"
                    except:
                        pass
                if "net::" in error_str.lower() or "NS_ERROR" in error_str:
                    return f"ERROR: 网络连接失败，请检查: 1) 是否需要代理 2) 网站是否可访问 3) DNS是否正常。\n详情: {error_str[:200]}"
                return f"ERROR: 导航失败: {error_str[:300]}"
        return self.browser_navigate(url, wait_until=wait_until, retries=3)

    # ==================== 浏览器登录 + 人工协助 ====================
    def browser_login(self, username, password, username_selector=None, password_selector=None, submit_selector=None):
        """
        在已打开的登录页面上自动填写账号密码并点击登录。
        - selector 留空则自动检测常见登录表单（input[type=email], input[name=username] 等）。
        - 登录后如果出现验证码，请调用 browser_wait_for_human 让用户手动完成。
        """
        if not hasattr(self, '_current_page'):
            return "ERROR: 未打开任何页面，请先调用 browser_navigate 打开登录页面"
        page = self._current_page
        try:
            # 自动检测选择器
            if not username_selector:
                for sel in ['input[type="email"]', 'input[name="username"]', 'input[name="account"]',
                            'input[name="mobile"]', 'input[placeholder*="手机"]', 'input[placeholder*="账号"]',
                            'input[placeholder*="用户名"]', 'input[placeholder*="邮箱"]', '#username', '#account']:
                    if page.query_selector(sel):
                        username_selector = sel
                        break
            if not password_selector:
                for sel in ['input[type="password"]', 'input[name="password"]', '#password']:
                    if page.query_selector(sel):
                        password_selector = sel
                        break
            if not submit_selector:
                for sel in ['button[type="submit"]', 'input[type="submit"]', 'button:has-text("登录")',
                            'button:has-text("登 录")', 'button:has-text("Sign in")', '.login-btn',
                            '#login-btn', '#submit', '.submit']:
                    if page.query_selector(sel):
                        submit_selector = sel
                        break

            if not username_selector or not password_selector:
                return "ERROR: 无法自动识别登录表单，请手动指定 username_selector / password_selector"
            if not submit_selector:
                return "ERROR: 无法自动识别登录按钮，请手动指定 submit_selector"

            page.fill(username_selector, '')
            page.type(username_selector, username, delay=random.randint(30, 100))
            page.fill(password_selector, '')
            page.type(password_selector, password, delay=random.randint(30, 100))
            page.click(submit_selector)
            page.wait_for_load_state("networkidle", timeout=15000)
            current_url = page.url
            return f"SUCCESS: 已填写账号密码并点击登录。当前URL: {current_url}\n\n⚠️ 如果需要验证码，请立即调用 browser_wait_for_human 等待用户手动完成验证码。"
        except Exception as e:
            return f"ERROR: 登录操作失败: {str(e)}"

    def browser_wait_for_human(self, seconds=120, message="请手动完成验证码操作"):
        """
        暂停自动化流程，弹出窗口等待用户手动操作（如输入验证码、滑块验证等）。
        用户点击"继续"后恢复执行。默认等待120秒。
        """
        import threading
        result = {"done": False}

        def _show_dialog():
            if self.gui is None:
                result["done"] = True
                return
            try:
                from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton
                from PyQt6.QtCore import Qt, QTimer
                def _show_on_main():
                    dlg = QDialog()
                    dlg.setWindowTitle("TGAI - 需要你的协助")
                    dlg.setFixedSize(420, 220)
                    dlg.setModal(True)
                    dlg.setWindowFlags(dlg.windowFlags() | Qt.WindowType.WindowStaysOnTopHint)
                    layout = QVBoxLayout(dlg)
                    layout.setContentsMargins(20, 20, 20, 20)
                    hint = QLabel(message)
                    hint.setWordWrap(True)
                    hint.setStyleSheet("font-size: 13pt;")
                    layout.addWidget(hint)
                    sub = QLabel(f"剩余等待时间：{seconds} 秒")
                    sub.setStyleSheet("color: gray; font-size: 10pt;")
                    layout.addWidget(sub)
                    btn = QPushButton("我已搞定，继续执行 ✅")
                    btn.setStyleSheet("background-color: #27ae60; color: white; font-weight: bold; padding: 8px 16px; font-size: 12pt;")
                    btn.clicked.connect(dlg.accept)
                    layout.addWidget(btn)

                    def _countdown(remaining):
                        if remaining <= 0:
                            dlg.reject()
                            return
                        sub.setText(f"剩余等待时间：{remaining} 秒")
                        QTimer.singleShot(1000, lambda: _countdown(remaining - 1))

                    QTimer.singleShot(1000, lambda: _countdown(seconds))
                    if dlg.exec() == QDialog.DialogCode.Accepted:
                        result["done"] = True
                    else:
                        result["done"] = True

                if hasattr(self.gui, 'schedule_on_main'):
                    self.gui.schedule_on_main(_show_on_main)
                else:
                    from PyQt6.QtCore import QTimer
                    QTimer.singleShot(0, _show_on_main)
            except Exception as e:
                print(f"[browser_wait_for_human] 弹窗失败: {e}")
                result["done"] = True

        t = threading.Thread(target=_show_dialog, daemon=True)
        t.start()
        timeout = max(seconds + 30, 30)
        t.join(timeout=timeout)
        if result["done"]:
            return "SUCCESS: 用户已完成手动操作，继续执行。"
        return "INFO: 等待超时，自动继续执行（用户可能未操作）。"

    # ==================== 内存优化工具 ====================
    def list_processes(self, top_n: int = 30):
        """列出当前正在运行的进程（按内存占用排序），标注可关闭/受保护"""
        from memory_optimizer import list_processes as _lp
        return _lp(top_n)

    def kill_process_by_name(self, name: str):
        """按进程名结束进程（受系统白名单保护，保护系统进程和 AI 自身）"""
        from memory_optimizer import kill_process_by_name as _kpn
        return _kpn(name)

    def kill_process_by_pid(self, pid: int):
        """按 PID 结束进程（受系统白名单保护，保护系统进程和 AI 自身）"""
        from memory_optimizer import kill_process_by_pid as _kpp
        return _kpp(pid)

    def optimize_memory(self):
        """执行系统内存优化：Python垃圾回收、释放工作集、统计临时文件"""
        from memory_optimizer import optimize_memory as _om
        return _om()

    def optimize_vram(self):
        """执行显存清理：CUDA显存释放、ONNX Runtime清理"""
        from memory_optimizer import optimize_vram as _ov
        return _ov()

    def full_memory_optimize(self):
        """执行完整内存+显存优化"""
        from memory_optimizer import full_optimize as _fo
        return _fo()

    # ==================== 桌面程序自动化工具 ====================
    def app_launch(self, exe_path: str, working_dir: str = None, args: str = None):
        """启动 Windows 程序（.exe 或 .lnk 快捷方式）"""
        from desktop_automation import launch_app
        return launch_app(exe_path, working_dir, args)

    def app_get_ui_tree(self, title: str = None, max_depth: int = 3):
        """获取桌面窗口的 UI 控件树，AI 通过它查看窗口中有哪些按钮、文本框等"""
        from desktop_automation import get_ui_tree
        return get_ui_tree(title, max_depth)

    def app_click_button(self, button_text: str = None, window_title: str = None, button_id: str = None):
        """通过按钮文字或 AutomationId 精准点击桌面程序中的按钮，无需坐标"""
        from desktop_automation import click_button
        return click_button(window_title=window_title, button_text=button_text, button_id=button_id)

    def app_type_text(self, text: str, window_title: str = None, field_name: str = None):
        """在当前焦点位置或指定输入框中输入文字"""
        from desktop_automation import type_into_app
        return type_into_app(text, window_title, field_name)

    def app_press_keys(self, keys: str):
        """发送键盘快捷键，如 'Ctrl+V'、'Alt+F4'、'Enter'"""
        from desktop_automation import press_keys
        return press_keys(keys)

    def app_screenshot(self, window_title: str = None, save_path: str = None):
        """截取指定窗口或全屏截图"""
        from desktop_automation import app_screenshot
        return app_screenshot(window_title, save_path)

    def app_set_clipboard(self, text: str):
        """设置剪贴板内容（配合 Ctrl+V 粘贴到不支持直接输入的字段）"""
        from desktop_automation import set_clipboard
        return set_clipboard(text)
